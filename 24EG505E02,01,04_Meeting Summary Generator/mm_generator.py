import groq
from docx import Document
from pydub import AudioSegment
import os
from dotenv import load_dotenv


class MeetingMinutesGenerator:
    def __init__(self) -> None:
        load_dotenv()
        self.client = groq.Groq()
        self.chat_model = "llama3-70b-8192"  # Model for text analysis
        self.transcription_model = "whisper-large-v3"  # Model for audio transcription
        self.output_folder = "output/"
        if not os.path.exists(self.output_folder):
            os.makedirs(self.output_folder, exist_ok=True)

    def transcribe_audio(self, audio_file_path):
        print("Getting the transcription for the audio...")
        try:
            # Load the audio file
            audio = AudioSegment.from_file(audio_file_path)
            
            # Split into 1-minute chunks (Groq has a file size limit)
            chunk_length_ms = 60 * 1000  # 1 minute
            chunks = []
            
            for i in range(0, len(audio), chunk_length_ms):
                chunk = audio[i:i + chunk_length_ms]
                chunk_path = os.path.join(self.output_folder, f"temp_chunk_{i}.wav")
                chunk.export(chunk_path, format="wav")
                chunks.append(chunk_path)
            
            # Transcribe each chunk
            full_transcription = []
            for chunk_path in chunks:
                with open(chunk_path, 'rb') as audio_file:
                    completion = self.client.audio.transcriptions.create(
                        model=self.transcription_model,
                        file=audio_file
                    )
                    full_transcription.append(completion.text)
                os.remove(chunk_path)  # Clean up the temporary file
            
            return " ".join(full_transcription)
            
        except Exception as e:
            print(f"Error during transcription: {str(e)}")
            return f"Error transcribing audio: {str(e)}"

    def meeting_minutes(self, transcription):
        return {
            'abstract_summary': self.__abstract_summary_extraction(transcription),
            'key_points': self.__key_points_extraction(transcription),
            'action_items': self.__action_item_extraction(transcription),
            'sentiment': self.__sentiment_analysis(transcription)
        }

    def __abstract_summary_extraction(self, transcription):
        print("Getting the summary ")
        prompt = f"Please provide a concise summary of the following text:\n\n{transcription}"
        
        completion = self.client.chat.completions.create(
            model=self.chat_model,
            messages=[{"role": "user", "content": prompt}]
        )
        return completion.choices[0].message.content

    def __key_points_extraction(self, transcription):
        print("Getting the key points ")
        prompt = f"Extract the main key points from this text as a bullet-point list:\n\n{transcription}"
        
        completion = self.client.chat.completions.create(
            model=self.chat_model,
            messages=[{"role": "user", "content": prompt}]
        )
        return completion.choices[0].message.content

    def __action_item_extraction(self, transcription):
        print("Getting the action items ")
        prompt = f"Extract all action items and tasks mentioned in this text as a bullet-point list:\n\n{transcription}"
        
        completion = self.client.chat.completions.create(
            model=self.chat_model,
            messages=[{"role": "user", "content": prompt}]
        )
        return completion.choices[0].message.content

    def __sentiment_analysis(self, transcription):
        print("Getting the sentiment analysis ")
        prompt = f"Analyze the overall sentiment and tone of this text. Consider aspects like positivity, negativity, formality, and emotional undertones:\n\n{transcription}"
        
        completion = self.client.chat.completions.create(
            model=self.chat_model,
            messages=[{"role": "user", "content": prompt}]
        )
        return completion.choices[0].message.content

    def save_as_docx(self, minutes, filename):
        print("Saving as .docx file...")
        output_path = os.path.join(self.output_folder, filename)
        
        # Make sure the output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # If file exists, try to remove it first
        if os.path.exists(output_path):
            try:
                os.remove(output_path)
            except Exception as e:
                print(f"Warning: Could not remove existing file: {str(e)}")
        
        try:
            doc = Document()
            for key, value in minutes.items():
                doc.add_heading(' '.join(word.capitalize() for word in key.split('_')), level=1)
                doc.add_paragraph(value)
                doc.add_paragraph()
            doc.save(output_path)
            print(f"Saved as {output_path}")
        except Exception as e:
            print(f"Error saving document: {str(e)}")
            # Try saving to a different location
            alt_path = os.path.join(os.path.expanduser("~"), "Desktop", filename)
            print(f"Attempting to save to alternative location: {alt_path}")
            doc.save(alt_path)
            print(f"Saved as {alt_path}")
